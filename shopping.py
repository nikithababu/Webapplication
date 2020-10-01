from tkinter import *
import docx
from docx import Document
from tkinter import ttk
import os
import openpyxl
import datetime


def login():
    global login_screen
    login_screen = Toplevel(main_screen)
    login_screen.title("Login")
    login_screen.geometry("300x250")
    Label(login_screen, text="Please enter details below to login").pack()
    Label(login_screen, text="").pack()

    global username_verify
    global password_verify

    username_verify = StringVar()
    password_verify = StringVar()

    global username_login_entry
    global password_login_entry

    Label(login_screen, text="Username * ").pack()
    username_login_entry = Entry(login_screen, textvariable=username_verify)
    username_login_entry.pack()
    Label(login_screen, text="").pack()
    Label(login_screen, text="Password * ").pack()
    password_login_entry = Entry(login_screen, textvariable=password_verify, show= '*')
    password_login_entry.pack()
    Label(login_screen, text="").pack()
    Button(login_screen, text="Login", width=10, height=1, command = login_verify).pack()

def login_verify():
    username1 = username_verify.get()
    password1 = password_verify.get()
    username_login_entry.delete(0, END)
    password_login_entry.delete(0, END)

    if (username1=="avanish") & (password1=="avanish@123"):
        login_sucess()
    else:
        Invalid_username_password()



def login_sucess():
    global login_success_screen
    login_success_screen = Toplevel(login_screen)
    login_success_screen.title("Success")
    login_success_screen.geometry("150x100")
    Label(login_success_screen, text="Login Success").pack()
    Button(login_success_screen, text="OK", command=delete_login_success).pack()


# Designing popup for login invalid password

def Invalid_username_password():
    global invalid
    invalid = Toplevel(login_screen)
    invalid.title("Wrong Entry")
    invalid.geometry("230x70")
    Label(invalid, text="Invalid Username or Password ").pack()
    Button(invalid, text="OK", command=delete_password_not_recognised).pack()


# Designing popup for user not found

def delete_login_success():
    login_success_screen.destroy()
    Data_change_write()

def Data_change_write():
    global login_screen
    login_screen = Toplevel(main_screen)
    login_screen.title("Login")
    login_screen.geometry("300x250")
    Label(login_screen, text="Please enter details below to login").pack()
    Label(login_screen, text="").pack()


def delete_password_not_recognised():
    invalid.destroy()

def customer_details():
    global customer_details
    customer_details = Toplevel(main_screen)
    customer_details.title("Customer Entry")
    customer_details.geometry("300x250+270+400")
    Label(customer_details, text="Please enter Customer Details for Purchase").pack()
    Label(customer_details, text="").pack()

    global name
    global contact_no
    global address

    name = StringVar()
    contact_no = StringVar()
    address = StringVar()

    global name_entry
    global contact_no_entry
    global address_entry

    Label(customer_details, text="Name : ").pack()
    name_entry = Entry(customer_details, textvariable=name)
    name_entry.pack()
    #Label(customer_details, text="").pack()
    Label(customer_details, text="Contact No:").pack()
    contact_no_entry = Entry(customer_details, textvariable=contact_no)
    contact_no_entry.pack()
    #Label(customer_details, text="").pack()
    Label(customer_details, text="Address:").pack()
    address_entry = Entry(customer_details, textvariable=address)
    address_entry.pack()
    Label(customer_details, text="").pack()

    Button(customer_details, text="Click for Shopping", width=20, height=2, command=Shopping_page).pack()

def Shopping_page():
    global shopping_page

    shopping_page = Toplevel(main_screen)
    shopping_page.title("Shopping")
    shopping_page.geometry("400x420+550+200")
    Label(shopping_page, text="Please enter No of Item required:").pack()
    # Label(shopping_page, text="").pack()
    n = StringVar()
    n1 = StringVar()
    Label(shopping_page, text=" Select Item you Want : Note Minimum Purchase 2 Items").pack()
    Label(shopping_page, text="").pack()

    Label(shopping_page, text="Item 1").pack()

    global itemlist

    itemlist = ttk.Combobox(shopping_page, width=35, textvariable=n)
    sample = openpyxl.load_workbook("electriclist.xlsx")
    ws = sample.active
    x = []
    for row in ws.rows:
        x.append(row[0].value)

    itemlist['values'] = x

    global Item_value_1

    itemlist.pack()
    frame = Frame(shopping_page)

    def increase():
        global value
        value = int(lbl_value["text"])
        lbl_value["text"] = f"{value + 1}"

    def decrease():
        value = int(lbl_value["text"])
        lbl_value["text"] = f"{value - 1}"

    Button(master=frame, text="-", width=6, height=1, command=decrease).pack(side=LEFT, fill=Y)
    lbl_value = Label(master=frame, text="0")
    lbl_value.pack(side=LEFT, fill=Y)
    Button(master=frame, text="+", width=6, height=1, command=increase).pack(side=LEFT, fill=Y)

    frame.pack()

    Label(shopping_page, text="Item 2").pack()
    global itemlist1

    itemlist1 = ttk.Combobox(shopping_page, width=35, textvariable=n1)

    sample = openpyxl.load_workbook("electriclist.xlsx")
    ws = sample.active
    x1 = []
    for row in ws.rows:
        x1.append(row[0].value)

    itemlist1['values'] = x1
    itemlist1.pack()

    frame1 = Frame(shopping_page)

    def increase1():
        global value1
        value1 = int(lbl_value1["text"])
        lbl_value1["text"] = f"{value1 + 1}"

    def decrease1():
        value1 = int(lbl_value1["text"])
        lbl_value1["text"] = f"{value1 - 1}"

    Button(master=frame1, text="-", width=6, height=1, command=decrease1).pack(side=LEFT, fill=Y)
    lbl_value1 = Label(master=frame1, text="0")
    lbl_value1.pack(side=LEFT, fill=Y)
    Button(master=frame1, text="+", width=6, height=1, command=increase1).pack(side=LEFT, fill=Y)

    frame1.pack()

    Button(shopping_page, text="Proceed", width=10, height=1, fg="green",command=save_data).pack()
    global data_save_msg
    data_save_msg = Label(shopping_page, text="", font=(12))
    data_save_msg.pack(side=BOTTOM)

    global item_Detail
    global item_Detail1
    item_Detail = Label(shopping_page, text="")
    item_Detail1 = Label(shopping_page, text="")

    global totalcost
    totalcost = Label(shopping_page, text="")
    itemlist.pack()

    data_save_msg=Label(shopping_page, text="", font=(12))
    data_save_msg.pack()
    item_Detail.pack()
    item_Detail1.pack()
    totalcost.pack()

    Button(shopping_page, text="Click to PRINT", font=" 12", width=20, height=2, command=printbill).pack()



def Confirmation_screen():
    global Confirmation_screen

    Confirmation_screen = Toplevel(main_screen)
    shopping_page.title("Confirmation_screen")
    shopping_page.geometry("200x200")
    Label(shopping_page, text="Confirm Product").pack()



def save_data():
    sample = openpyxl.load_workbook("Customer_purchase_details.xlsx")
    ws = sample.active

    item_purchased = itemlist.get()
    item_purchased1 = itemlist1.get()
    name1 = name.get()
    contact_no1 = contact_no.get()
    address1= address.get()

    name_entry.delete(0, END)
    contact_no_entry.delete(0, END)
    address_entry.delete(0, END)
    itemlist.delete(0, END)
    itemlist1.delete(0, END)

    sample1 = openpyxl.load_workbook("electriclist.xlsx")
    ws1 = sample1.active

    for row in ws1.rows:
        # print(row[1].value)
        if (row[0].value == item_purchased):
            Item_value_1 = row[1].value
        if (row[0].value == item_purchased1):
            Item_value_1_1 = row[1].value

    global p
    global p1

    p = float(Item_value_1) * (value + 1)
    p1= float(Item_value_1_1) * (value1 + 1)
    n = "Price of " + item_purchased + " is " + str(Item_value_1) + ". Total cost of Item 1: " + str(p)
    m = "Price of " + item_purchased1 + " is " + str(Item_value_1_1) + ". Total cost of Item 2 : " + str(p1)


    current_row = ws.max_row
    ws.cell(row=current_row + 1, column=1).value = name1
    ws.cell(row=current_row + 1, column=2).value = contact_no1
    ws.cell(row=current_row + 1, column=3).value = address1
    ws.cell(row=current_row + 1, column=4).value = datetime.datetime.now()
    ws.cell(row=current_row + 1, column=5).value = item_purchased
    ws.cell(row=current_row + 1, column=6).value = value+1
    ws.cell(row=current_row + 1, column=7).value = Item_value_1
    ws.cell(row=current_row + 1, column=8).value = p
    ws.cell(row=current_row + 1, column=9).value = item_purchased1
    ws.cell(row=current_row + 1, column=10).value = value1+1
    ws.cell(row=current_row + 1, column=11).value = Item_value_1_1
    ws.cell(row=current_row + 1, column=12).value = p1
    ws.cell(row=current_row + 1, column=13).value = p+p1
    sample.save("Customer_purchase_details.xlsx")

    t=p+p1

    item_Detail["text"] = n
    item_Detail1["text"] = m
    data_save_msg["text"] = "Data Stored"
    totalcost["text"]= "Total Cost: "+ str(t)
    customer_details.destroy()

def printbill():
    doc= Document()

    doc.add_heading('Electronic Shop', 0)
    x = 'xyz, state, country'
    p = doc.add_paragraph(x)
    p.add_run('- XXXXXX').bold = True


    sample = openpyxl.load_workbook("Customer_purchase_details.xlsx")
    ws = sample.active
    x1 = []
    abc=ws.max_row

    for row in ws.columns:
        x1.append(row[1].value)

    name = x1[0]
    mobile = x1[1]
    address = x1[2]

    doc.add_paragraph('Customer Name: ' + name, style='List Bullet')
    doc.add_paragraph('Contact No:' + mobile, style='List Bullet')
    doc.add_paragraph('Address:' + address, style='List Bullet')


    records = [ [1,x1[4], x1[5],x1[6],x1[7]],
                [2,x1[8], x1[9],x1[10],x1[11] ]]

    print(records)
    print(x1)


    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Serial No'
    hdr_cells[1].text = 'Item Name'
    hdr_cells[2].text = 'Quantity'
    hdr_cells[3].text = 'Price'
    hdr_cells[4].text = 'Total price'

    for ser, item, quna, price, totalprice in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(ser)
        row_cells[1].text = str(item)
        row_cells[2].text = str(quna)
        row_cells[3].text = str(price)
        row_cells[4].text = str(totalprice)

    totalvalue = x1[12]
    totalvalue1 = str(totalvalue)
    doc.add_paragraph('Total cost: ' + totalvalue1)
    doc.add_paragraph(" some quotation")

    doc.add_page_break()
    doc.save('demo.docx')


def main_account_screen():
    global main_screen
    main_screen = Tk()
    main_screen.geometry("400x350")
    main_screen.title("Shop Software")
    Label(text="Welcome to Shop", bg="orange", width="400", height="2", font=("times", 18)).pack()
    Label(text="").pack()
    Button(text="Shop Owner Login", height="3", width="30", command = login).pack()
    Label(text="").pack()
    Button(text="Customer Purchase", height="3", width="30", command=customer_details).pack()
    Label(text="").pack()
    Button(text="Statistics of Product", height="3", width="30").pack()

    main_screen.mainloop()


main_account_screen()
